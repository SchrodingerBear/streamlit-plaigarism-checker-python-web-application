import streamlit as st
import plotly.express as px
from utils import (
    add_custom_css,
    check_plagiarism,
    get_top_repeated_words,
    get_limitations
)
try:
    import docx
    import PyPDF2
    import io
except ImportError as e:
    st.error(f"""
    Missing document processing packages. Please install them using:
    pip install python-docx PyPDF2
    Error: {str(e)}
    """)
    st.stop()

st.set_page_config(layout="centered")
add_custom_css()
st.title("FLAG-AI: PLAGIARISM DETECTOR")

st.markdown("<div style='display: flex; justify-content: center;'>", unsafe_allow_html=True)
text_area = st.text_area("Enter or paste text here", height=200, max_chars=25000, key="text_input", label_visibility="collapsed")


if text_area and len(text_area) >= 25000:
    st.warning("⚠️ You have reached the maximum character limit of 25,000 characters.")

def read_docx(file) -> str:
    """
    Read and extract text from a DOCX file
    Args:
        file: A file-like object containing the DOCX file
    Returns:
        str: Extracted text from the document
    """
    try:
        doc = docx.Document(file)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip(): 
                full_text.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Error reading DOCX file: {str(e)}")
        return ""

def read_pdf(file) -> str:
    """
    Read and extract text from a PDF file
    Args:
        file: A file-like object containing the PDF file
    Returns:
        str: Extracted text from the document
    """
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        full_text = []
        
        for page in pdf_reader.pages[:10]:
            text = page.extract_text()
            if text.strip(): 
                full_text.append(text)
        
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Error reading PDF file: {str(e)}")
        return ""

def process_uploaded_file(uploaded_file) -> str:
    """
    Process an uploaded file based on its type
    Args:
        uploaded_file: Streamlit's UploadedFile object
    Returns:
        str: Extracted text from the document
    """
    try:
        file_size = uploaded_file.size
        max_size = 5 * 1024 * 1024 
        
        if file_size > max_size:
            st.error("File size exceeds 5MB limit. Please upload a smaller file.")
            return ""
        
        if uploaded_file.type == "text/plain":
            return uploaded_file.read().decode("utf-8")
        
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return read_docx(uploaded_file)
        
        elif uploaded_file.type == "application/pdf":
            return read_pdf(uploaded_file)
        
        else:
            st.error(f"Unsupported file type: {uploaded_file.type}")
            return ""
            
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return ""

uploaded_file = st.file_uploader(
    "Upload a file", 
    type=["txt", "docx", "pdf"],
    accept_multiple_files=False,
    help="Limit 5MB per file • TXT, DOCX, PDF"
)

if uploaded_file is not None:
    text_content = process_uploaded_file(uploaded_file)
    if text_content:
        st.text_area("Extracted Text", text_content, height=200)

st.markdown("<div style='display: flex; justify-content: center;'>", unsafe_allow_html=True)
analyze_button = st.button("Analyze")
st.markdown("</div>", unsafe_allow_html=True)

api_key = 'AIzaSyD7pDUCZfKOibhM5fTlY2yl0TLWjHgUO_g'
cse_id = '84e326e702f3a4031'

if analyze_button:
    text_to_check = text_area
    if uploaded_file:
        text_to_check = text_content

    if not text_to_check.strip():
        st.warning("Please enter some text or upload a file to check for plagiarism.")
    else:
        with st.spinner("Checking for plagiarism... This will only take a few seconds"):
            plagiarism_results, highlighted_text, plagiarism_score, originality_score, message = check_plagiarism(text_to_check, api_key, cse_id)

        if message:
            st.info(message)

        col1, col2 = st.columns(2)

        with col1:
            st.subheader("Your Input Text")
            st.markdown(f'<div class="input-text">{highlighted_text}</div>', unsafe_allow_html=True)

        with col2:
            st.subheader("Detection Score")
            if plagiarism_score > 30:
                st.markdown(f'<div class="detection-result">Text Analysis Result: Plagiarised Content Detected<br><span class="detection-score">Detection Score: {plagiarism_score:.1f}%</span></div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="detection-result">Text Analysis Result: No Significant Plagiarism Detected<br><span class="detection-score">Detection Score: {plagiarism_score:.1f}%</span></div>', unsafe_allow_html=True)

            st.markdown(f'<div class="disclaimer">{get_limitations()}</div>', unsafe_allow_html=True)

        col3, col4 = st.columns(2)

        with col3:
            st.subheader("Basic Insights")
            st.markdown('Top 5 Most Repeated Words', unsafe_allow_html=True)
            
            top_words = get_top_repeated_words(text_to_check)
            fig = px.bar(top_words, x='word', y='count', title='')
            fig.update_layout(
                xaxis_title="Words",
                yaxis_title="Counts",
                margin=dict(l=20, r=20, t=40, b=20),
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
            )   
            st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': False})

        with col4:  
            st.subheader("Potential Source Websites")
            if plagiarism_results:
                for result in plagiarism_results:
                    st.markdown(f"""
                    🔗 Source: [{result['url']}]({result['url']})
                    """)
                    if 'matched_parts' in result and result['matched_parts']:
                        matched_parts = result['matched_parts'][:2]
                        if matched_parts:
                            st.markdown("🔍 Matched Content:")
                            for part in matched_parts:
                                st.markdown(f"- {part[:100]}...")
            else:
                st.info("No matching sources found.")